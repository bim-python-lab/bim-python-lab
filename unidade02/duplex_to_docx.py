# Programação Open BIM com Python, (c) PUCRS, 2025

import ifcopenshell
from docx import Document

model = ifcopenshell.open("duplex.ifc")
types = ["IfcWall", "IfcDoor", "IfcWindow"]

doc = Document()
doc.add_heading("Relatório IFC Modelo Duplex", 0)

for t in types:
    elements = model.by_type(t)
    doc.add_heading(f"{t}: {len(elements)} elementos", level=1)
    for e in elements:
        name = getattr(e, "Name", "")
        gid = e.GlobalId
        doc.add_paragraph(f"{name} - {gid}")

doc.save("duplex_docx.docx")